import os
import re
from typing import List, Dict

from docx import Document


def fetch_denodo_metadata(config_file: str) -> List[Dict[str, str]]:
    """Parse a configuration file and extract metadata information.

    Currently, this function only extracts the configured heap memory amount.
    Extend the parsing logic as needed for other metrics.
    """
    metadata: List[Dict[str, str]] = []

    with open(config_file, 'r', encoding='utf-8', errors='ignore') as f:
        content = f.read()

    heap_value = None

    # Look for typical JVM heap settings like '-Xmx2G'
    match = re.search(r'-Xmx([\d]+[mMgG])', content)
    if not match:
        # Also check for patterns like 'heap memory: 2048m'
        match = re.search(r'heap\s*memory\s*[:=]?\s*([\d.]+\s*[kKmMgG][bB]?)', content)

    if match:
        heap_value = match.group(1)

    metadata.append({
        'database': 'configuration',
        'view': 'heap_memory',
        'description': heap_value or 'Unknown'
    })

    return metadata


def generate_word_report(metadata: List[Dict[str, str]], output_path: str) -> None:
    """Generate a Word report from metadata."""
    document = Document()
    document.add_heading("Denodo Metadata Report", level=1)

    for item in metadata:
        document.add_heading(f"{item['database']}.{item['view']}", level=2)
        document.add_paragraph(item.get("description", "No description"))

    document.save(output_path)


if __name__ == "__main__":
    config_file = os.getenv("DENODO_CONFIG_FILE", "config.txt")
    output = "denodo_report.docx"

    metadata = fetch_denodo_metadata(config_file)
    generate_word_report(metadata, output)
    print(f"Report saved to {output}")
